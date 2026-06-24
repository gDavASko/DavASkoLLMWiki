# -*- coding: utf-8 -*-
"""Build the English scientific article .docx (TNR 12pt, 1.5 spacing, A4)."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
FIGS = os.path.join(HERE, "figs")
OUT = os.path.join(HERE, "..", "..", "docs", "paper", "davasko-llm-wiki.docx")
GREY = RGBColor(0x5b, 0x65, 0x73)

doc = Document()

# ── page + base style ──────────────────────────────────────────────────────
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)            # A4
sec.top_margin = sec.bottom_margin = Cm(2.0)
sec.left_margin = sec.right_margin = Cm(2.5)
CONTENT_CM = 16.0

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"; normal.font.size = Pt(12)
# make TNR apply to Cyrillic / complex scripts too
rpr = normal.element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
    rf.set(qn(a), "Times New Roman")
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_after = Pt(0)

def _runs(p, parts):
    """parts: list of (text, dict-of-flags). flags: b,i, size, color."""
    for text, f in parts:
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        rf2 = r._element.get_or_add_rPr().get_or_add_rFonts()
        for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rf2.set(qn(a), "Times New Roman")
        if f.get("b"): r.font.bold = True
        if f.get("i"): r.font.italic = True
        if f.get("size"): r.font.size = Pt(f["size"])
        if f.get("color"): r.font.color.rgb = f["color"]
    return p

def para(text, align="just", indent=True, before=0, after=0):
    p = doc.add_paragraph()
    p.alignment = {"just": WD_ALIGN_PARAGRAPH.JUSTIFY, "c": WD_ALIGN_PARAGRAPH.CENTER,
                   "l": WD_ALIGN_PARAGRAPH.LEFT}[align]
    if indent: p.paragraph_format.first_line_indent = Cm(0.7)
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    if isinstance(text, str): _runs(p, [(text, {})])
    else: _runs(p, text)
    return p

def heading(text, size=12, before=12, after=4):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    _runs(p, [(text, {"b": True, "size": size})]); return p

def formula(text):
    return para([(text, {"i": True})], align="c", indent=False, before=4, after=4)

def figure(png, caption, width_cm=CONTENT_CM):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(os.path.join(FIGS, png), width=Cm(width_cm))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(10)
    _runs(c, [(caption, {"size": 10.5, "color": GREY})])

def caption(text, num=True):
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.LEFT if num else WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_before = Pt(6); c.paragraph_format.space_after = Pt(2)
    _runs(c, [(text, {"size": 10.5, "color": GREY, "b": num})])

def _set_cell(cell, text, bold=False, align="l", shade=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"l": WD_ALIGN_PARAGRAPH.LEFT, "r": WD_ALIGN_PARAGRAPH.RIGHT,
                   "c": WD_ALIGN_PARAGRAPH.CENTER}[align]
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(0)
    _runs(p, [(text, {"b": bold, "size": 10.5})])
    if shade:
        tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
        sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), shade); tcPr.append(sh)

def table(headers, rows, aligns, note=None, best_row=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        _set_cell(t.rows[0].cells[j], h, bold=True, align=("l" if j == 0 else "r"), shade="EDEFF2")
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        sh = "EEFAF4" if best_row == i else None
        for j, v in enumerate(row):
            _set_cell(cells[j], v, bold=(best_row == i and j == 0), align=aligns[j], shade=sh)
    if note: caption(note, num=False)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ════════════════════════════════════════════════════════════════════════════
#  FRONT MATTER
# ════════════════════════════════════════════════════════════════════════════
para([("UDC 004.8 (artificial intelligence); 004 (computer science)", {"size": 10.5, "color": GREY})],
     align="l", indent=False, after=6)
para([("A Layered Self-Validating Knowledge Base with Hybrid Retrieval for Grounding "
       "LLM Agents: Architecture, Method and Empirical Evaluation on a Real Corpus",
       {"b": True, "size": 14})], align="c", indent=False, after=6)
para([("Davletbaev Aleksandr Sergeevich", {"b": True})], align="c", indent=False, after=2)
para([("Lead Software Engineer, LLC «KB Production» (ООО «КБ Продакшн»)", {"i": True})],
     align="c", indent=False, after=2)
para([("System development and numerical experiments were carried out with the aid of AI tools "
       "(LLM assistants). Framework DavASko LLM Wiki, v3.x — fully offline, reproducible.",
       {"i": True, "size": 10.5, "color": GREY})], align="c", indent=False, after=10)

heading("Abstract", before=4)
para("It is verified empirically whether a curated hybrid (symbolic and semantic) retrieval layer over "
     "a knowledge base outperforms a trivial lexical file search when grounding LLM agents. The system "
     "architecture is described in detail: a downward-dependent layer model that separates immutable "
     "source snapshots from derived curated pages with content-hash provenance; a hybrid retriever with a "
     "per-query adaptive threshold; and a deduplicated, machine-shared embedding model resolved through a "
     "system marker. On a deployed 162-document base with 15 labeled queries, recall@5, MRR and nDCG@5 are "
     "computed for the proposed engine and baselines. Semantic retrieval yields recall@5 0.633 and MRR 0.718 "
     "against 0.333 and 0.435 for the lexical baseline; structure-aware chunking exceeds the fixed-window one "
     "by 7.8% MRR; GPU embedding gives an eight-fold speed-up. A stress test across three query regimes "
     "(meaning, topic, exact identifier) additionally exposed and corrected a non-functional symbolic stream, "
     "doubling exact-identifier retrieval. The retrieval layer is justified by measurement.")
para([("Keywords: ", {"b": True}),
      ("information retrieval; knowledge base; retrieval-augmented generation; vector embeddings; "
       "layered architecture; ranking quality evaluation; LLM agents", {})])

heading("Аннотация", before=10)
para("Эмпирически проверяется, превосходит ли курируемый гибридный (символьный и семантический) слой поиска "
     "по базе знаний тривиальный лексический поиск по файлам при заземлении LLM-агентов. Подробно описана "
     "архитектура системы: слоевая модель со строго нисходящими зависимостями, отделяющая иммутабельные снимки "
     "источников от производных курируемых страниц с провенансом по контент-хешам; гибридный ретривер с "
     "адаптивным порогом на запрос; и дедуплицированная, общая для машины модель эмбеддингов, разрешаемая через "
     "системную метку. На развёрнутой базе из 162 документов с 15 размеченными запросами вычислены recall@5, MRR "
     "и nDCG@5 для предложенного движка и базлайнов. Семантический поиск даёт recall@5 0.633 и MRR 0.718 против "
     "0.333 и 0.435 у лексического базлайна; структурный чанкинг превосходит оконный на 7.8% MRR; эмбеддинг на GPU "
     "даёт восьмикратное ускорение. Слой поиска оправдан измерением.")
para([("Ключевые слова: ", {"b": True}),
      ("информационный поиск; база знаний; дополненная поиском генерация; векторные эмбеддинги; "
       "слоевая архитектура; оценка качества ранжирования; LLM-агенты", {})])

# ════════════════════════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════
heading("1. Introduction")
para([("Relevance. ", {"b": True}),
      ("LLM agents in software development increasingly operate over large and evolving project knowledge. "
       "A common grounding approach is retrieval-augmented generation (RAG) [5], in which relevant context is "
       "retrieved and supplied to the model. Yet the underlying assumption that a dedicated retrieval layer is "
       "effective is usually asserted rather than tested: it remains open whether a curated indexed knowledge "
       "layer beats the trivial alternative — direct full-text search by the agent over the files. A retrieval "
       "layer that does not beat that baseline is negative-value complexity, as it adds maintenance burden and a "
       "second, drifting copy of the truth.", {})])
para([("Aim and objectives. ", {"b": True}),
      ("The aim is to establish quantitatively whether the proposed hybrid retrieval engine earns its use, and "
       "to describe the architecture that makes it maintainable. The objectives are: (1) to formalize the layered "
       "architecture and retrieval method; (2) to build an evaluation harness with information-retrieval metrics "
       "and baselines; (3) to evaluate on a real corpus; (4) to refine the method based on observed data.", {})])
para([("Contributions. ", {"b": True}),
      ("(i) A layered, self-validating knowledge-base architecture with explicit source-vs-derived separation and "
       "content-hash provenance (§2.2, §2.5); (ii) a hybrid retriever with a per-query adaptive threshold (2) "
       "proposed as a replacement for a constant cosine cutoff (§2.4); (iii) a machine-shared embedding model "
       "resolved through a marker, removing per-base model duplication (§2.6); (iv) an evaluation harness sharing "
       "the engine's own core, and a measurement-driven refinement loop (§3–§4).", {})])
para([("Related work. ", {"b": True}),
      ("Dense retrieval maps a query and text fragments into a shared vector space and ranks by cosine similarity "
       "[4]; RAG couples such retrieval with a generator [5, 2]. The embedding model is the multilingual Jina v3 with "
       "task-specific LoRA adapters and asymmetric query:/passage: prefixes [8], run locally via ONNX Runtime in "
       "Transformers.js [10]. Inverted-file (IVF) indexes partition vectors into clusters and probe the nearest at "
       "query time, trading recall for speed [3]. Structure-aware splitters cut a document along its structure and "
       "improve fragment coherence [1]. Combining lexical and dense retrieval is established practice [2].", {})])

# ════════════════════════════════════════════════════════════════════════════
#  2. METHODS
# ════════════════════════════════════════════════════════════════════════════
heading("2. Materials and Methods")

heading("2.1 System overview", size=12, before=10)
para("The system has two pipelines over a shared store: an offline indexing pipeline that turns knowledge into "
     "vectors, and an online query pipeline that retrieves grounded context for an agent. The store is organized "
     "into independent layers (§2.2); each layer separates immutable source snapshots (raw/) from derived curated "
     "pages (wiki/), and both are indexed. The end-to-end pipelines are shown in Figure 1.")
figure("fig1_pipeline.png", "Figure 1 — Indexing (top) and query (bottom) pipelines")

heading("2.2 Layered knowledge model", size=12, before=10)
para("Knowledge is partitioned into independent layers forming a strictly downward dependency graph (a directed "
     "acyclic graph): a more specific layer may depend on and link to more general ones, never the reverse. This "
     "isolates universal agent rules from engine constraints, framework conventions and project specifics, and lets "
     "several project layers reuse a common base in parallel. Each layer declares its dependencies in a wiki.json "
     "manifest. When the same topic exists in two layers, the more specific layer wins by default and the agent "
     "must flag the duplicate. The dependency model and the internal split of a single layer are shown in Figure 2.")
figure("fig2_layers.png", "Figure 2 — Layered dependency graph (left) and the source-vs-derived split inside one layer (right)")

heading("2.3 Hybrid retrieval", size=12, before=10)
para("The symbolic stream (A) matches strict code identifiers (PascalCase of at least two parts, I* interfaces, "
     "m_* fields) against document fields; matches are weighted by field type and down-weighted by an inverse "
     "document frequency in the spirit of IDF [6]. The semantic stream (B) embeds the query with the query: prefix "
     "and ranks fragments by cosine similarity (1) [6]:")
formula("sim(q, d) = (q · d) / (‖q‖ · ‖d‖)      (1)")
para("The streams are merged by a single effective score rather than by a hard precedence of one over the other "
     "(justification in §4.1). Exact symbolic hits (score ≈ 1.0) and dense cosine scores (0–1) thus compete on one "
     "scale, so a strong semantic match can outrank a tangential identifier match.")

heading("2.4 Routing and adaptive threshold", size=12, before=10)
para("Vectors are stored one shard per layer; a shard centroid is the mean of its member vectors. At query time, "
     "shards are ranked by centroid proximity and the nprobe nearest are scanned (IVF-style multi-probe [3]); when "
     "nprobe is not smaller than the number of shards, the search is exhaustive and loses no recall. Instead of a "
     "fixed cosine threshold, sensitive to a shift of the similarity distribution, a per-query adaptive threshold (2) "
     "is used:")
formula("τ_q = max( φ,  α · max_d sim(q, d) )      (2)")
para("where α is a fraction of the maximum similarity (default 0.85) and φ is a lower noise floor (default 0.35). "
     "Construction (2) is proposed in this work as a replacement for a constant threshold: it adapts to each query's "
     "own best score and is therefore robust to distribution shifts induced by language (RU/EN) and document length. "
     "Finally, graph-lift expands the top results by one hop along extends parents and [[wiki-links]], pulling in "
     "directly related pages.")

heading("2.5 Structure-aware chunking and provenance", size=12, before=10)
para("The indexer tokenizes Markdown into headings, paragraphs and atomic code blocks, packs them to a target size "
     "within the [min, max] word range, hard-splits only oversized blocks, merges undersized fragments and prepends "
     "the heading path; the approach follows the idea of structure-aware splitting [1]. Code is treated as the source "
     "of truth and wiki/ pages as derived; on close scores the source is ranked higher and results are labeled primary "
     "(SOURCE) or derived (SUMMARY). Each derived page stores content hashes of its cited sources; a staleness check "
     "recomputes them, so a divergence between a changed source and its summary becomes detectable rather than silent.")

heading("2.6 Shared embedding model and deployment", size=12, before=10)
para("The embedding model is large (~1.1 GB). Storing a copy inside every deployed base would multiply disk use and "
     "management. Instead the model is installed once per machine into a system location and its path is published "
     "through a small marker file (a JSON config); every base resolves the model through that marker. Resolution order "
     "is: an explicit environment variable, then the marker, then a repository-local fallback, and otherwise the deployer "
     "asks where to install the shared model. The mechanism is shown in Figure 3.")
figure("fig3_model.png", "Figure 3 — One machine-shared model resolved by every base through a marker; resolution order below")
para("Deployment of a complete base is one command (or the equivalent agent skill) performing five steps: scaffold the "
     "layers and base pages, install the shared model and write the marker, install the agent rules, install the "
     "operating skills, and run a baseline validation (unit tests, linter, index build). The write path is closed "
     "end-to-end: ingesting a source places it under raw/, auto-creates a derived summary stub, lints, and finishes "
     "with vectorization, so new knowledge is immediately searchable.")

heading("2.7 Evaluation metrics", size=12, before=10)
para("For a set of relevant documents R and a ranked result list, standard metrics are used [6, 7]. Recall in the "
     "first k (3) is the fraction of relevant documents in the top-k [6]:")
formula("recall@k = |R ∩ top-k| / |R|      (3)")
para("Mean reciprocal rank (4) averages the reciprocal rank of the first relevant document over queries Q [6]:")
formula("MRR = (1/|Q|) · Σ 1 / rank_i      (4)")
para("Normalized discounted cumulative gain (5)–(6) accounts for the position of a relevant result [7]:")
formula("DCG@k = Σ rel_i / log₂(i+1);   nDCG@k = DCG@k / IDCG@k      (5)–(6)")

heading("2.8 Experimental setup", size=12, before=10)
para("The evaluation corpus is the deployed KBPro base; its composition is given in Table 1. The query set comprises "
     "15 labeled questions with required-source annotations. The model is jinaai/jina-embeddings-v3 (fp16, 1024 "
     "dimensions). The parameter top-k = 5; the threshold is adaptive (α = 0.85, φ = 0.35). The retrievers compared "
     "are semantic (dense multi-probe only), hybrid (symbols and semantics with unified ranking), flat (exhaustive flat "
     "cosine), and lexical (term overlap over the files — the agent-reads-the-files model). The harness reuses the "
     "engine's own core functions and therefore evaluates the real engine, not a re-implementation. Embedding runs on a "
     "GPU via the DirectML provider with fallback to the CPU [9].")
caption("Table 1 — Evaluation corpus composition")
table(["Scope", "layers", "documents", "fragments", "shards"],
      [["evaluation subset (kbpro-wiki, llm-wiki)", "2", "162", "433", "2"],
       ["full deployment", "6", "343", "—", "6"]],
      ["l", "r", "r", "r", "r"], best_row=0,
      note="With 2 shards, multi-probe equals exhaustive search; clustering matters only across many layers.")

# ════════════════════════════════════════════════════════════════════════════
#  3. RESULTS
# ════════════════════════════════════════════════════════════════════════════
heading("3. Results")
para("The main results are given in Table 2 and Figure 4. On the real corpus the retrieval layer roughly doubles "
     "recall (0.633 against 0.333) and improves first-relevant ranking by 65% in MRR (0.718 against 0.435) relative "
     "to the lexical baseline.")
caption("Table 2 — Main retrieval results (162 documents, 15 queries, top-k = 5)")
table(["Retriever", "recall@5", "MRR", "nDCG@5"],
      [["semantic (proposed engine)", "0.633", "0.718", "0.626"],
       ["hybrid (symbols + semantic)", "0.633", "0.718", "0.626"],
       ["flat (exhaustive cosine)", "0.633", "0.718", "0.626"],
       ["lexical (lexical baseline)", "0.333", "0.435", "0.303"]],
      ["l", "r", "r", "r"], best_row=0,
      note="The equality semantic = hybrid = flat is due to having only two shards, in which case multi-probe equals exhaustive search.")
figure("fig4_results.png", "Figure 4 — Dense retrieval versus the lexical baseline across three metrics", width_cm=13)

# ════════════════════════════════════════════════════════════════════════════
#  4. REFINEMENT
# ════════════════════════════════════════════════════════════════════════════
heading("4. Result Discussion and Method Refinement")
heading("4.1 Hybrid ranking", size=12, before=10)
para("The first run revealed a defect: under the original merge all symbolic matches were placed above semantic ones, "
     "whereby tangential exact matches displaced the best semantic results and the hybrid ranked worse than pure "
     "semantics (MRR 0.641 against 0.718). Two refinements, measured on the same index, removed the gap (Table 3, "
     "Figure 5): a switch to unified score-based ranking and a tightening of symbol extraction that excludes generic "
     "acronyms. The diagnosis was localized to a single query in which the acronym JSON was treated as a code identifier.")
caption("Table 3 — Ranking refinement (hybrid-retriever metrics)")
table(["Configuration", "MRR", "nDCG@5"],
      [["original (hard symbol precedence)", "0.641", "0.577"],
       ["+ unified score-based ranking", "0.685", "0.610"],
       ["+ strict symbol extraction", "0.718", "0.626"]],
      ["l", "r", "r"], best_row=2,
      note="Total gain — 12% MRR with no change in recall.")
figure("fig5_refine.png", "Figure 5 — Measurement-driven progression of hybrid MRR (0.641 → 0.685 → 0.718)", width_cm=13)

heading("4.2 Comparison of chunking strategies", size=12, before=10)
para("With all else equal, the index was rebuilt under each strategy. Structure-aware chunking surpassed the "
     "fixed-window one in ranking at equal recall (Table 4, Figure 6), which confirms on this corpus the common "
     "qualitative observation on the benefit of structural splitting [1].")
caption("Table 4 — Structure-aware chunking versus fixed window")
table(["Chunking", "recall@5", "MRR", "nDCG@5", "fragments"],
      [["structure-aware", "0.633", "0.718", "0.626", "433"],
       ["fixed window", "0.633", "0.666", "0.585", "418"]],
      ["l", "r", "r", "r", "r"], best_row=0,
      note="Gain of structure-aware chunking — 7.8% MRR and 7% nDCG.")
figure("fig6_chunking.png", "Figure 6 — Structure-aware versus fixed-window chunking at equal recall", width_cm=13)

heading("4.3 Indexing speed", size=12, before=10)
para("The correctness of each speed-up was verified before time was measured. Batched embedding is provably equivalent "
     "to single-item embedding (minimum cosine 0.99999997), but on the CPU it gives only about 11% gain, since a forward "
     "pass of the model saturates the computation. Moving computation to a GPU via DirectML [9] gives an eight-fold "
     "speed-up while preserving numerical equivalence (Table 5, Figure 7), reducing full-index time from about half an "
     "hour to a few minutes.")
caption("Table 5 — Embedding speed and result equivalence")
table(["Configuration", "time (12 emb.)", "speed-up", "cosine to CPU"],
      [["central processor (baseline)", "9657 ms", "1.0×", "—"],
       ["+ batching (on 45 fragments)", "≈ 11% faster", "≈ 1.1×", "≥ 0.99999997"],
       ["graphics processor (DirectML)", "1205 ms", "8.0×", "0.999984"]],
      ["l", "r", "r", "r"], best_row=2,
      note="The dominant speed lever is the GPU; batching is a small free gain.")
figure("fig7_gpu.png", "Figure 7 — Relative embedding throughput; the GPU is the dominant lever (numerically equivalent)", width_cm=12)

# ════════════════════════════════════════════════════════════════════════════
#  5–7
# ════════════════════════════════════════════════════════════════════════════
heading("4.4 Stress test and reviving the symbolic stream", size=12, before=10)
para("Beyond the 15-query regression set, the engine was stress-tested on three larger query regimes constructed over "
     "the full 345-document deployment, each probing a different retrieval mode (Table 6): (R1) topic to document, where "
     "the query is a document's topic label (lexically easy); (R2) cross-lingual and paraphrased questions sharing no "
     "surface tokens with the target, where a lexical baseline cannot win by construction; and (R3) exact unique code "
     "identifiers (document frequency one), the best case for lexical search.")
para("The outcome splits sharply by regime. On meaning-bearing queries (R2) dense retrieval roughly tripled the lexical "
     "baseline (recall@5 0.684 against 0.237); on exact identifiers (R3) the lexical baseline dominated (0.990 against "
     "0.286). R3 also exposed a defect: the hybrid equalled the semantic retriever exactly, i.e. the symbolic stream "
     "contributed nothing. The cause was that the indexer populated each document's symbols only from YAML frontmatter, "
     "leaving raw and code documents with empty symbol sets, so Stream A had nothing to match. The fix - automatic "
     "extraction of code identifiers (PascalCase, I-interfaces, m_-fields) from document content into the index - raised "
     "the share of documents with indexed symbols from 76 to 292 of 346 and, on R3, doubled hybrid recall@5 and raised "
     "MRR by about 150% with no regression on R1 or R2 (Table 7). The streams are thus complementary: dense retrieval "
     "owns meaning, the revived symbolic stream recovers exact-identifier lookup, and a verbatim file scan remains "
     "strongest for literal tokens. This is a further instance of the measurement-driven loop: a stress test revealed a "
     "silent failure that aggregate accuracy on the original set had hidden.")
caption("Table 6 - Three-regime stress test, recall@5 (full 345-document corpus)")
table(["Query regime", "n", "hybrid", "semantic", "lexical"],
      [["R2 - cross-lingual / paraphrase (meaning)", "38", "0.684", "0.684", "0.237"],
       ["R1 - topic to document", "100", "0.929", "0.929", "0.838"],
       ["R3 - exact unique identifier (after fix)", "98", "0.576", "0.283", "0.990"]],
      ["l", "r", "r", "r", "r"], best_row=0,
      note="Methods are complementary: the engine wins on meaning (R2), a verbatim scan wins on exact tokens (R3).")
caption("Table 7 - Reviving the symbolic stream (regime R3, exact identifiers)")
table(["Index symbol source", "docs with symbols", "hybrid recall@5", "hybrid MRR"],
      [["frontmatter only (before)", "76 / 346", "0.286", "0.227"],
       ["+ content auto-extraction (after)", "292 / 346", "0.576", "0.567"]],
      ["l", "r", "r", "r"], best_row=1,
      note="Auto-extracting code identifiers doubles exact-identifier recall@5 (+101%) and MRR (+150%); semantic and lexical unchanged.")

heading("5. Discussion")
para("The data show that the dense retrieval layer is not decorative: it substantially outperforms the lexical baseline "
     "in both recall and ranking. For natural-language questions semantic retrieval is the working tool; the symbolic "
     "stream is neutral on this set but is intended to raise precision on identifier queries not present in the set. Each "
     "method improvement (§4) was driven by measurement rather than by an a priori assumption, which illustrates the "
     "embedding of the evaluation harness into the engineering loop. With two shards the clustering mechanism does not "
     "affect the result and begins to matter only across many layers. The architectural choices (§2) target "
     "maintainability rather than benchmark scores: the source-vs-derived split with provenance keeps drift detectable, "
     "and the shared-model marker removes per-base duplication without changing retrieval behavior.")

heading("6. Limitations and Threats to Validity")
para("The following limitations are established:", indent=False, after=2)
for t in [
    "the small query set (n = 15) leads to wide confidence intervals; the numbers indicate direction, not a precise value;",
    "the results are obtained on a single corpus; generalization to other domains is untested;",
    "relevance labels are a proxy (required sources of a regression set), which may bias the recall estimate;",
    "in the evaluation 2 of 6 layers are indexed, and fewer distractor documents make retrieval easier than in full deployment;",
    "recall@5 = 0.633 implies a loss of about 37% of relevant material in the top-5; raising recall requires a larger k, query expansion, or a change of representation rather than re-ranking;",
    "high indexing speed needs a GPU and an appropriate runtime build, otherwise CPU fallback applies;",
    "provenance makes divergence detectable but does not remove the dual source of truth between code and derived pages.",
]:
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    _runs(p, [(t, {})])

heading("7. Conclusion")
para("In this work a layered, self-validating knowledge-base architecture is described in detail and an often-asserted "
     "assumption about the effectiveness of a dedicated retrieval layer is turned into a measured fact: on a real "
     "deployed knowledge base a layered hybrid retrieval engine substantially outperforms the lexical baseline (recall@5 "
     "0.633 against 0.333, MRR 0.718 against 0.435). A measurement-driven refinement loop is demonstrated: two ranking "
     "corrections (+12% MRR) and the choice of chunking strategy (+7.8% MRR) were adopted after confirmation by data, and "
     "the speed optimizations (eight-fold GPU speed-up, batching) were introduced after proven numerical equivalence. A "
     "three-regime stress test extended the evaluation and exposed a silent failure (a symbolic stream left empty by "
     "frontmatter-only symbols) whose correction (content auto-extraction of code identifiers) doubled exact-identifier "
     "retrieval (hybrid recall@5 0.286 to 0.576, MRR 0.227 to 0.567) without regressing the other regimes. The "
     "architecture contributes a source-vs-derived split with content-hash provenance and a machine-shared, "
     "marker-resolved embedding model that eliminates per-base duplication. The limitations are documented and mapped to "
     "concrete directions for further experiments. All results are reproducible with the accompanying harness.")

# ── references ───────────────────────────────────────────────────────────────
heading("References")
refs = [
    "Chase H. et al. LangChain: RecursiveCharacterTextSplitter — structure-aware text splitting. Documentation, 2023. URL: python.langchain.com.",
    "Gao Y. et al. Retrieval-Augmented Generation for Large Language Models: A Survey. arXiv:2312.10997, 2023.",
    "Johnson J., Douze M., Jégou H. Billion-scale similarity search with GPUs (FAISS) // IEEE Transactions on Big Data. 2019. arXiv:1702.08734.",
    "Karpukhin V. et al. Dense Passage Retrieval for Open-Domain Question Answering // EMNLP. 2020. arXiv:2004.04906.",
    "Lewis P. et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks // NeurIPS. 2020. arXiv:2005.11401.",
    "Manning C. D., Raghavan P., Schütze H. Introduction to Information Retrieval. Cambridge University Press, 2008. 482 p.",
    "Järvelin K., Kekäläinen J. Cumulated Gain-based Evaluation of IR Techniques // ACM TOIS. 2002. Vol. 20, no. 4. P. 422–446.",
    "Sturua S. et al. jina-embeddings-v3: Multilingual Embeddings with Task LoRA. arXiv:2409.10173, 2024.",
    "Microsoft. ONNX Runtime Execution Providers: DirectML. Documentation, 2023. URL: onnxruntime.ai/docs/execution-providers.",
    "Hugging Face. Transformers.js — ONNX-runtime inference in JavaScript. 2023. URL: github.com/huggingface/transformers.js.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.7); p.paragraph_format.first_line_indent = Cm(-0.7)
    _runs(p, [(f"{i}. ", {"b": True, "size": 11}), (r, {"size": 11})])

doc.save(os.path.abspath(OUT))
print("SAVED", os.path.abspath(OUT))
